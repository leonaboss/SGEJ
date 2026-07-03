import os
import hashlib
import base64
import re
import logging
from io import BytesIO
from cryptography.fernet import Fernet
from django.conf import settings
from django.core.exceptions import ImproperlyConfigured
from django.utils import timezone
from docx import Document as DocxDocument
from apps.documentos.models import Documento
import qrcode

logger = logging.getLogger(__name__)


class DocumentoSecurizadoService:
    _master_fernet_cache = None

    @staticmethod
    def _get_master_fernet():
        if DocumentoSecurizadoService._master_fernet_cache is not None:
            return DocumentoSecurizadoService._master_fernet_cache
        key = settings.FERNET_MASTER_KEY
        if isinstance(key, str):
            key = key.encode('utf-8')
        try:
            DocumentoSecurizadoService._master_fernet_cache = Fernet(key)
        except Exception as e:
            raise ImproperlyConfigured(
                "FERNET_MASTER_KEY inválida o no configurada en .env. "
                "Debe ser una clave Fernet de 32 bytes en base64. "
                "Genere una con: python -c 'from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())'"
            ) from e
        return DocumentoSecurizadoService._master_fernet_cache

    @staticmethod
    def procesar_y_cifrar_archivo(archivo_plano, expediente_obj, usuario_creador):
        sha256_context = hashlib.sha256()
        for chunk in archivo_plano.chunks():
            sha256_context.update(chunk)
        hash_sha256 = sha256_context.hexdigest()

        file_key = Fernet.generate_key()
        file_fernet = Fernet(file_key)

        archivo_plano.seek(0)
        contenido_binario = archivo_plano.read()
        contenido_cifrado = file_fernet.encrypt(contenido_binario)

        master_fernet = DocumentoSecurizadoService._get_master_fernet()
        encrypted_key = master_fernet.encrypt(file_key)
        encrypted_key_b64 = base64.urlsafe_b64encode(encrypted_key).decode('utf-8')

        nombre_original = archivo_plano.name
        nombre_cifrado_uuid = f"{hash_sha256}.enc"
        ruta_destino = os.path.join(settings.MEDIA_ROOT, 'securized_docs', nombre_cifrado_uuid)

        os.makedirs(os.path.dirname(ruta_destino), exist_ok=True)
        with open(ruta_destino, 'wb') as f:
            f.write(contenido_cifrado)

        qr_data = f"SGIJ-EXPEDIENTE:{expediente_obj.numero_expediente}|HASH:{hash_sha256}"
        qr = qrcode.QRCode(version=1, box_size=10, border=5)
        qr.add_data(qr_data)
        qr.make(fit=True)
        img_qr = qr.make_image(fill_color="black", back_color="white")

        qr_filename = f"QR_{hash_sha256}.png"
        qr_path = os.path.join(settings.MEDIA_ROOT, 'qrs', qr_filename)
        os.makedirs(os.path.dirname(qr_path), exist_ok=True)
        img_qr.save(qr_path)

        documento_instancia = Documento.objects.create(
            expediente=expediente_obj,
            nombre_original=nombre_original,
            nombre_cifrado=nombre_cifrado_uuid,
            tipo_mime=archivo_plano.content_type,
            hash_sha256=hash_sha256,
            iv_cifrado=encrypted_key_b64,
            qr_code_content=qr_data,
            version=1,
            created_by=usuario_creador
        )

        return documento_instancia

    @staticmethod
    def descifrar_archivo(documento):
        ruta = os.path.join(settings.MEDIA_ROOT, 'securized_docs', documento.nombre_cifrado)
        if not os.path.exists(ruta):
            raise FileNotFoundError("El archivo físico no se encuentra en el servidor.")
        encrypted_key_b64 = documento.iv_cifrado
        encrypted_key = base64.urlsafe_b64decode(encrypted_key_b64.encode('utf-8'))
        master_fernet = DocumentoSecurizadoService._get_master_fernet()
        file_key = master_fernet.decrypt(encrypted_key)
        file_fernet = Fernet(file_key)
        with open(ruta, 'rb') as f:
            contenido_cifrado = f.read()
        contenido_descifrado = file_fernet.decrypt(contenido_cifrado)
        return contenido_descifrado

    @staticmethod
    def aplicar_marca_agua(contenido_pdf, texto_marca):
        try:
            from reportlab.pdfgen import canvas
            from pypdf import PdfReader, PdfWriter
            reader = PdfReader(BytesIO(contenido_pdf))
            writer = PdfWriter()
            for page in reader.pages:
                pw = float(page.mediabox.width)
                ph = float(page.mediabox.height)
                packet = BytesIO()
                c = canvas.Canvas(packet, pagesize=(pw, ph))
                c.saveState()
                c.setFillColorRGB(0.5, 0.5, 0.5, alpha=0.25)
                c.setFont("Helvetica", 24)
                c.translate(pw / 2, ph / 2)
                c.rotate(45)
                c.drawCentredString(0, 0, texto_marca)
                c.restoreState()
                c.save()
                packet.seek(0)
                watermark = PdfReader(packet)
                if watermark.pages:
                    page.merge_translated_page(watermark.pages[0], 0, 0, over=True)
                writer.add_page(page)
            output = BytesIO()
            writer.write(output)
            return output.getvalue()
        except Exception as e:
            logger.warning("Error al aplicar marca de agua: %s", e)
            return contenido_pdf


class PlantillaService:

    @staticmethod
    def obtener_variables(plantilla):
        try:
            text = plantilla.archivo_plantilla.read().decode('utf-8', errors='ignore')
            plantilla.archivo_plantilla.seek(0)
            return sorted(set(re.findall(r'\{\{(\w+)\}\}', text)))
        except Exception as e:
            logger.warning("Error al obtener variables de plantilla: %s", e)
            return []

    @staticmethod
    def generar_documento(plantilla, valores, usuario, expediente=None):
        template_doc = DocxDocument(plantilla.archivo_plantilla)
        for paragraph in template_doc.paragraphs:
            for key, val in valores.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(val))
        for table in template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in valores.items():
                        if f'{{{{{key}}}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{{key}}}}}', str(val))
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)
        contenido = output.read()

        sha256_hash = hashlib.sha256(contenido).hexdigest()
        file_key = Fernet.generate_key()
        file_fernet = Fernet(file_key)
        contenido_cifrado = file_fernet.encrypt(contenido)

        master_fernet = DocumentoSecurizadoService._get_master_fernet()
        encrypted_key = master_fernet.encrypt(file_key)
        encrypted_key_b64 = base64.urlsafe_b64encode(encrypted_key).decode('utf-8')

        nombre_doc = f"{plantilla.nombre}_{usuario.usuario}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.docx"
        nombre_cifrado_uuid = f"{sha256_hash}.enc"
        ruta = os.path.join(settings.MEDIA_ROOT, 'securized_docs', nombre_cifrado_uuid)
        os.makedirs(os.path.dirname(ruta), exist_ok=True)
        with open(ruta, 'wb') as f:
            f.write(contenido_cifrado)

        doc = Documento.objects.create(
            expediente=expediente,
            nombre_original=nombre_doc,
            nombre_cifrado=nombre_cifrado_uuid,
            tipo_mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            hash_sha256=sha256_hash,
            iv_cifrado=encrypted_key_b64,
            es_plantilla=False,
            created_by=usuario,
        )
        return doc
